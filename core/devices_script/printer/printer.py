import win32print as printer
import win32api
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm
import sys
import os
from win32printing import Printer


PRINTER_ERROR_STATES = (
    printer.PRINTER_STATUS_NO_TONER, #262144
    printer.PRINTER_STATUS_NOT_AVAILABLE, #4096
    printer.PRINTER_STATUS_OFFLINE, #128
    printer.PRINTER_STATUS_OUT_OF_MEMORY, #2097152
    printer.PRINTER_STATUS_OUTPUT_BIN_FULL, #2048
    printer.PRINTER_STATUS_PAGE_PUNT, #524288
    printer.PRINTER_STATUS_PAPER_JAM, #8
    printer.PRINTER_STATUS_PAPER_OUT, #16
    printer.PRINTER_STATUS_PAPER_PROBLEM, #64
)

def codeToString(code):
    match code:
        case printer.PRINTER_STATUS_NO_TONER:
            return 'Out of toner'
        case printer.PRINTER_STATUS_NOT_AVAILABLE:
            return 'Printer not available'
        case printer.PRINTER_STATUS_OFFLINE:
            return 'Printer offline'
        case printer.PRINTER_STATUS_OUT_OF_MEMORY:
            return 'Printer out of memory'
        case printer.PRINTER_STATUS_OUTPUT_BIN_FULL:
            return 'Printer bin is full'
        case printer.PRINTER_STATUS_PAGE_PUNT:
            return 'Printer page punt'
        case printer.PRINTER_STATUS_PAPER_JAM:
            return 'Printer paper jam'
        case printer.PRINTER_STATUS_PAPER_OUT:
            return 'Printer out of paper'  
        case printer.PRINTER_STATUS_PAPER_PROBLEM:
            return 'Printer paper problem'    
        

def printer_errorneous_state(prn, error_states=PRINTER_ERROR_STATES):
    prn_opts = printer.GetPrinter(prn)
    status_opts = prn_opts[18]
    for error_state in error_states:
        if status_opts & error_state:
            return error_state
    return 0

def startPrint():
    printer_name = "BK-C310(U) 1" # or get_printer_names()[0]
    prn = printer.OpenPrinter(printer_name)
    error = printer_errorneous_state(prn)
    if error:
        print("error occurred: ", codeToString(error))
    else:
        #  Do the real work
        print( win32api.ShellExecute(0, "print", 'demo.docx', None,  "./devices_script/printer",  0) )
    printer.ClosePrinter(prn)
    if not error:
        print('Printed')

def writeToFile(name,room):
    try:
        with open('./devices_script/printer/input.txt', 'r') as f:
            str = f.read()
    except:
        return False

    str = str.replace('$NAME$', name)
    str = str.replace('$ROOM$', room)

    try:
        with open('./devices_script/printer/output.txt', 'w') as f:
            f.write(str)
    except:
        return False
    print('File written')
    return True

def writeImage():
    print('ok')

    document = Document()

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Mm(30.4)
    p.paragraph_format.right_indent = Mm(14.4)
    
    r   = p.add_run()
    # r.add_text(str)
    r.add_text('Benvenuti')
    r.add_break()
    r.add_text('Hilton Hotels & Resorts')
    r.add_break()
    r.add_text('Nome: Lorenzo')
    r.add_break()
    r.add_text('Stanza: 034')
    r.add_break()

    # r.add_text(' do you like it?')
    r.add_picture('./devices_script/printer/qrcode.png',width=Inches(1.0))

    document.save('./devices_script/printer/demo.docx')

# print ('Argument List:', str(sys.argv))
# print ('Number of arguments:', len(sys.argv), 'arguments.')
writeImage()
startPrint()
if len(sys.argv) == 3:
    result = writeToFile(sys.argv[1], sys.argv[2])
else:
    result = writeToFile(sys.argv[1]+' '+sys.argv[2], sys.argv[3])

if result:
    startPrint()
else:
    print('error writing to file')